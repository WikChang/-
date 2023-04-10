import os
import pandas as pd
from docx import Document

def add_data_to_paragraph(doc, field, value):
    for paragraph in doc.paragraphs:
        if field in paragraph.text:
            paragraph.text = paragraph.text.replace(field, f"{field} {value}")

def create_word_from_template(template_path, excel_path, output_path):
        #确保输出文件夹存在
    if not os.path.exists(output_path):
        os.makedirs(output_path)
    # 读取Excel文件
    df = pd.read_excel(excel_path)
    print(df)
    # 对数据进行分组，根据复试地点和简略时间进行分组
    grouped_data = df.groupby(["复试地点", "简略时间"])

    # 遍历分组数据
    for group, group_data in grouped_data:
        location, time = group

        # 读取模板文件
        doc = Document(template_path)

        # 获取表格
        table = doc.tables[0]

        # 添加招生专业信息
        majors = group_data["复试专业名称"].unique()
        major_text = " & ".join([f"{major}" for major in majors])
        add_data_to_paragraph(doc, "招生专业：", major_text)

        # 添加复试时间信息
        add_data_to_paragraph(doc, "复试时间：", f"{time}")

        # 从第二行开始添加数据（假设第一行为标题行）
        current_row_index = 1

        # 将分组数据添加到表格中
        for _, row_data in group_data.iterrows():
            row = table.rows[current_row_index].cells
            row[0].text = str(row_data["序号"])
            row[1].text = row_data["考生姓名"]
            current_row_index += 1

        # 生成文件名
        filename = f"各科平均分汇总表-{location}-{time}.docx"
        output_path = os.path.join(output_path, filename)

        # 保存生成的Word文件
        doc.save(output_path)

# 调用函数，指定模板、Excel和输出文件夹路径
template_path = "./input_file/各科平均分汇总表.docx"
excel_path = "./input_file/面试名单.xlsx"
output_path = "./output_file/各科平均分汇总表"

create_word_from_template(template_path, excel_path, output_path)